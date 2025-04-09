import docx2txt
import PyPDF2
import google.generativeai as genai
import os
import re
import time
from dotenv import load_dotenv
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config import GEMINI_MODEL
import logging

# 🤖 Load Environment Variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

# 📝 Logging Setup
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# =====================================
# 🤖 Gemini API Call
# =====================================
def call_gemini(prompt, retries=3):
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        response = model.generate_content(prompt)
        logging.debug("🔍 Gemini Response: %s", response.text)
        return response.text.strip() if response and hasattr(response, "text") else "Error: No response from Gemini."
    except Exception as e:
        if retries > 0:
            time.sleep(2)
            return call_gemini(prompt, retries - 1)
        logging.error("Error calling Gemini API: %s", str(e))
        return f"Error calling Gemini API: {str(e)}"

# =====================================
# 📄 Resume File Text Extraction
# =====================================
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
    except Exception as e:
        logging.error("Error extracting text from PDF: %s", str(e))
        return f"Error extracting text from PDF: {str(e)}"
    return text.strip()

def extract_text_from_docx(docx_path):
    try:
        return docx2txt.process(docx_path).strip()
    except Exception as e:
        logging.error("Error extracting text from DOCX: %s", str(e))
        return f"Error extracting text from DOCX: {str(e)}"

def process_resume(file_path):
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        return "Unsupported file format. Please upload a PDF or DOCX file."

# =====================================
# 📝 Feedback Formatting
# =====================================
def format_feedback(raw_feedback):
    formatted_feedback = raw_feedback.strip()
    formatted_feedback = re.sub(r'\*\*(.*?)\*\*', r'\1', formatted_feedback)  # Keep bold text without **
    formatted_feedback = re.sub(r'\n{2,}', '\n\n', formatted_feedback)
    return formatted_feedback

# =====================================
# 🌟 Skill Extraction from Gemini Feedback
# =====================================
def extract_skills_from_feedback(feedback, section):
    try:
        pattern = rf"{section}:\s*\n((?:-.*(?:\n|$))*)"
        matches = re.search(pattern, feedback, re.IGNORECASE)
        if matches:
            skills_text = matches.group(1).strip()
            return [line.strip("- ").strip() for line in skills_text.splitlines() if line.strip().startswith("-")]
        return []
    except Exception as e:
        logging.error("Error extracting %s: %s", section, str(e))
        return []

# =====================================
# 🤖 Match Resume with Job Description
# =====================================
def match_resume_with_job(resume_text, job_description):
    try:
        prompt = f"""
You are an expert career advisor and resume evaluator.

Your task is to analyze the candidate's resume in comparison with the given job description and return a structured report.

⚠️ Strictly follow the format below. Do NOT add anything else.

FORMAT TO FOLLOW:
Match Score: [X]%

Missing Skills:
- List technical and soft skills that are relevant to the job but not present in the resume.

Strengths:
- Highlight skills, projects, or experiences from the resume that match well with the job description.

Areas for Improvement:
- Provide specific suggestions to align the resume better with the job role, such as skills to add, formatting tips, or ways to reword existing content.

Name: [Name of the resume holder]
Email: [Email of the resume holder]
Phone: [Phone number of the resume holder]

📄 Resume:
{resume_text}

📝 Job Description:
{job_description}
"""
        response = call_gemini(prompt)
        logging.debug("🔍 Raw Gemini Response: %s", response)

        if "Match Score" not in response:
            logging.error("Gemini response missing Match Score")
            return {
                "match_percentage": 0,
                "feedback": "⚠️ Gemini failed to return structured feedback. Please try again.\n\nRaw Response:\n" + response,
                "matched_skills": [],
                "missing_skills": [],
                "Name": "N/A",
                "email": "N/A",
                "phone": "N/A"
            }

        # Split response and process
        lines = response.splitlines()
        feedback_lines = []
        name, email, phone = "N/A", "N/A", "N/A"
        match_percentage = 0

        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith("Match Score:"):
                match_percentage = int(re.search(r"(\d+)%", line).group(1)) if re.search(r"(\d+)%", line) else 0
                feedback_lines.append(line)
            elif line.startswith("Name:") and name == "N/A":
                name = line.replace("Name:", "").strip()
            elif line.startswith("Email:") and email == "N/A":
                email = line.replace("Email:", "").strip()
            elif line.startswith("Phone:") and phone == "N/A":
                phone = line.replace("Phone:", "").strip()
            elif any(line.startswith(s) for s in ["Missing Skills:", "Strengths:", "Areas for Improvement:"]):
                feedback_lines.append(line)
                i += 1
                while i < len(lines) and (lines[i].strip().startswith("-") or lines[i].strip() == ""):
                    if lines[i].strip().startswith("-"):
                        feedback_lines.append(lines[i].strip())
                    i += 1
                continue
            i += 1

        # Reconstruct feedback
        feedback_text = "\n\n".join(feedback_lines)
        logging.debug("Constructed feedback: %s", feedback_text)

        matched_skills = extract_skills_from_feedback(feedback_text, "Strengths")
        missing_skills = extract_skills_from_feedback(feedback_text, "Missing Skills")

        logging.debug("Matched Skills: %s", matched_skills)
        logging.debug("Missing Skills: %s", missing_skills)

        return {
            "Name": name,
            "email": email,
            "phone": phone,
            "match_percentage": match_percentage,
            "feedback": format_feedback(feedback_text),
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
        }

    except Exception as e:
        logging.error("Error processing Gemini response: %s", str(e))
        return {
            "match_percentage": 0,
            "feedback": f"⚠️ Error processing Gemini response: {str(e)}",
            "matched_skills": [],
            "missing_skills": [],
            "Name": "N/A",
            "email": "N/A",
            "phone": "N/A"
        }

# =====================================
# 📝 Generate Improved Resume
# =====================================

def generate_improved_resume(resume_text, analysis_result, output_path):
    improved_text = call_gemini(f"""
Based on the following resume and analysis feedback, generate an improved version:

Resume:
{resume_text}

Analysis Feedback:
{analysis_result}

Format it as a single-column layout using this structure:
📌 **Professional Summary**
🛠️ **Skills**
💼 **Experience**
🎓 **Education**
📜 **Certifications** (if any)

Use bullet points, emoji markers, and keep the tone professional.
Quantify achievements and improve readability. Keep it stylish, clean, and modern.
""")

    # Replace double asterisks for headings and single for bullet points
    improved_text = re.sub(r"\*\*(.+?)\*\*", r"<<HEADING:\1>>", improved_text)
    improved_text = re.sub(r"\*(.+?)\*", r"- \1", improved_text)

    doc = Document()

    # Add title
    title = doc.add_heading("\U0001F4BC Improved Resume", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for paragraph in improved_text.split("\n\n"):
        lines = paragraph.strip().split("\n")
        if not lines:
            continue

        for line in lines:
            line = line.strip()

            # Solid heading
            if line.startswith("<<HEADING:") and line.endswith(">>"):
                heading_text = line.replace("<<HEADING:", "").replace(">>", "").strip()
                para = doc.add_paragraph()
                run = para.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 102, 204)
                para.paragraph_format.space_before = Pt(10)

            # Bullet point
            elif line.startswith("-") or line.startswith("•"):
                doc.add_paragraph(line.lstrip("-• ").strip(), style="List Bullet")

            # Regular line
            elif line:
                para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)

    # Add footer note
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        "\n\U0001F4A1 Note: Focus on improving the highlighted target areas and quantify achievements when updating your resume."
    )
    note_run.italic = True
    note_para.paragraph_format.space_before = Pt(10)

    doc.save(output_path)